import json
import re

from docx import Document


# 使用示例
doc_path = 'C:\\Users\\Administrator\\Downloads\\中华成语典故大全.docx'  # 替换为你的Word文档路径
# sections = read_doc_and_split_by_headings(doc_path)
doc = Document(doc_path)
# 遍历文档中的每个段落
list=[]
for para in doc.paragraphs:
    list.append(para.text)
    print(para.text)

# 初始化一个变量来存储当前页的内容
current_page_content = []

# 遍历文档中的每个段落
for para in doc.paragraphs:
    # 检查段落样式是否为分页符
    if para.style.name == 'PAGE':
        # 如果是分页符，打印当前页的内容并重置当前页内容存储
        print('\n'.join(current_page_content))
        current_page_content = []
    else:
        # 如果不是分页符，将段落文本添加到当前页内容
        current_page_content.append(para.text)

# 打印最后一页的内容（如果没有分页符，则打印整个文档的内容）

def split_by_empty_triples(lst):
    # 初始化结果列表
    result = []
    # 初始化一个临时列表来存储当前片段的元素
    temp = []

    # 遍历输入列表
    for item in lst:
        # 如果当前元素不是空字符串，则添加到临时列表
        if item:
            temp.append(item)
        else:
            # 如果当前元素是空字符串，检查是否连续出现三个
            if len(temp) > 0 and len(temp) % 2 == 0:
                # 如果临时列表的长度是3的倍数，将临时列表添加到结果列表
                result.append(temp)
                # 清空临时列表以准备下一个片段
                temp = []
            else:
                # 如果不是三个连续的空字符串，将非空元素添加到临时列表
                temp.append(item)

    # 检查最后一个片段是否需要添加到结果列表
    if temp:
        result.append(temp)

    return result

result=split_by_empty_triples(current_page_content)
result_json=[]
for item in result[2:]:
    item_tet="\n".join(item)
    if "释义" in item_tet:
        name, detail = item_tet.split("释义", 1)
        # 将分割后的结果添加到结果列表中
        result_json.append({
            "name": name,
            "detail": detail
        })

# 如果需要将JSON保存到文件中
with open('中华成语典故大全_清洗.json', 'w', encoding='utf-8') as json_file:
    json.dump(result_json, json_file, ensure_ascii=False, indent=4)


